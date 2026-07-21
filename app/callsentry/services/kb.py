"""Knowledge base: document ingestion, chunking, and pgvector retrieval."""

from __future__ import annotations

import io
import re
import uuid
from dataclasses import dataclass

import structlog
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession

from callsentry.models import KBChunk, KBDocument
from callsentry.services.embeddings import get_embeddings

log = structlog.get_logger(__name__)

CHUNK_CHARS = 900
CHUNK_OVERLAP = 150


class UnsupportedDocument(ValueError):
    pass


def extract_text(filename: str, data: bytes, content_type: str = "") -> str:
    """Pull plain text out of PDF / DOCX / TXT / MD."""
    lower = filename.lower()

    if lower.endswith(".pdf") or content_type == "application/pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)

    if lower.endswith(".docx"):
        import docx

        document = docx.Document(io.BytesIO(data))
        return "\n".join(p.text for p in document.paragraphs)

    if lower.endswith((".txt", ".md", ".markdown")) or content_type.startswith("text/"):
        return data.decode("utf-8", errors="replace")

    raise UnsupportedDocument(f"unsupported file type: {filename} ({content_type})")


def chunk_text(text: str, *, size: int = CHUNK_CHARS, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Split on paragraph boundaries, packing up to `size` chars per chunk.

    Overlap carries the tail of the previous chunk forward so an answer that
    straddles a boundary is still retrievable from one chunk.
    """
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    current = ""

    for para in paragraphs:
        if len(current) + len(para) + 2 <= size:
            current = f"{current}\n\n{para}" if current else para
            continue
        if current:
            chunks.append(current)
            current = (current[-overlap:] + "\n\n" + para) if overlap else para
        else:
            # A single paragraph longer than the window: hard-split it.
            for i in range(0, len(para), size - overlap):
                chunks.append(para[i : i + size])
            current = ""

    if current:
        chunks.append(current)
    return chunks


@dataclass
class Retrieved:
    chunk_text: str
    filename: str
    score: float
    document_id: uuid.UUID


async def index_document(
    session: AsyncSession,
    *,
    business_id: uuid.UUID,
    filename: str,
    data: bytes,
    content_type: str = "",
) -> tuple[KBDocument, bool]:
    """Extract, chunk, embed, and store. Returns (document, embeddings_ok)."""
    text = extract_text(filename, data, content_type)
    if not text.strip():
        raise UnsupportedDocument(f"no extractable text in {filename}")

    pieces = chunk_text(text)
    embedding = await get_embeddings().embed(pieces)

    document = KBDocument(
        business_id=business_id,
        filename=filename,
        content_type=content_type or "text/plain",
        content=text,
        chunk_count=len(pieces),
    )
    session.add(document)
    await session.flush()

    for i, piece in enumerate(pieces):
        vector = embedding.vectors[i] if i < len(embedding.vectors) else None
        session.add(
            KBChunk(
                document_id=document.id,
                business_id=business_id,
                chunk_index=i,
                chunk_text=piece,
                embedding=vector,
            )
        )

    await session.flush()
    log.info(
        "kb.indexed",
        filename=filename,
        chunks=len(pieces),
        provider=embedding.provider,
        degraded=embedding.degraded,
    )
    return document, not embedding.degraded


async def search(
    session: AsyncSession,
    *,
    business_id: uuid.UUID,
    query: str,
    limit: int = 4,
) -> list[Retrieved]:
    """Cosine-similarity retrieval scoped to one business."""
    vector, result = await get_embeddings().embed_one(query)
    if result.degraded:
        # Zero vector would rank arbitrarily; return nothing so the caller
        # escalates rather than answering from noise.
        return []

    distance = KBChunk.embedding.cosine_distance(vector)
    rows = (
        await session.execute(
            select(KBChunk, KBDocument.filename, distance.label("distance"))
            .join(KBDocument, KBChunk.document_id == KBDocument.id)
            .where(KBChunk.business_id == business_id, KBChunk.embedding.isnot(None))
            .order_by(distance)
            .limit(limit)
        )
    ).all()

    return [
        Retrieved(
            chunk_text=chunk.chunk_text,
            filename=filename,
            # pgvector cosine_distance is in [0,2]; map to a [0,1] similarity.
            score=max(0.0, 1.0 - float(dist)),
            document_id=chunk.document_id,
        )
        for chunk, filename, dist in rows
    ]


async def delete_document(
    session: AsyncSession, *, business_id: uuid.UUID, document_id: uuid.UUID
) -> bool:
    document = await session.get(KBDocument, document_id)
    if document is None or document.business_id != business_id:
        return False
    await session.execute(delete(KBChunk).where(KBChunk.document_id == document_id))
    await session.delete(document)
    return True
